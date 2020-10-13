from __future__ import absolute_import
from pdf2docx import parse
from docx import Document
import urllib.request
from flask import Flask, jsonify
from flask_sqlalchemy import SQLAlchemy
import os, re, datetime
import atexit
from sqlalchemy.ext.declarative import DeclarativeMeta
import json
from flask import request
import yaml, requests
from apscheduler.schedulers.background import BackgroundScheduler
import hashlib
from database.db import db
from database.base import *

dirname = os.path.dirname(os.path.realpath(__file__))
fileloc = dirname + "/nadconfig.yaml"
with open(fileloc) as yamlfile:
    config = yaml.load(yamlfile, Loader=yaml.FullLoader)
if config == None:
    raise FileNotFoundError("Yaml file does not exist")


def debug(*args):
    print(f"[DEBUG] [{datetime.datetime.now()}] {' '.join(args)}")


debug("App started")

app = Flask(__name__)


db = db(config)


def getday(url):
    datestr = re.search("_in_obvestila_(.+).pdf", url, re.IGNORECASE).group(1)
    print(datestr)
    date_time_obj = datetime.datetime.strptime(datestr, "%d._%m._%Y")
    return date_time_obj.isoformat().split("T")[0]


# headers of pdf
nad = [
    "ODSOTNI",
    "UČITELJ/ICA",
    "URA",
    "RAZRED",
    "UČILNICA",
    "NADOMEŠČA",
    "PREDMET",
    "OPOMBA",
]
menur = [
    "RAZRED URA",
    "RAZRED URA",
    "UČITELJ/ICA",
    "PREDMETA UČILNICA",
    "PREDMETA UČILNICA",
    "OPOMBA",
]
menuc = [
    "RAZRED",
    "URA",
    "UČITELJ/ICA",
    "PREDMET",
    "IZ",
    "UČILNICE",
    "V",
    "UČILNICO",
    "OPOMBA",
]


class PdfFile:
    def __init__(self, url, pdftype="nd"):
        self.url = url
        self.pdftype = pdftype
        self.pdfloc = dirname + "/tmp/tempdf.pdf"
        self.docloc = dirname + "/tmp/tempdoc.docx"

    def download(self):
        u = urllib.request.urlopen(self.url)
        with open(self.pdfloc, "wb") as fl:
            fls = u.read()
            fl.write(fls)
            h = hashlib.sha256(fls).hexdigest()
        return str(h)

    def parsepdf(self):
        self.download()
        parse(self.pdfloc, self.docloc)
        ttable = []
        document = Document(self.docloc)
        for table in document.tables:
            tabl = []
            for row in table.rows:
                rrow = []
                for cell in row.cells:
                    ccell = ""
                    for para in cell.paragraphs:
                        ccell += para.text.replace("\n", " ") + " "
                    rrow.append(" ".join(ccell.split()))
                tabl.append(rrow)
            ttable.append(tabl)
        print(ttable)

        return self.parsenad(ttable)

    def parsenad(self, ttable):
        nadobjects = []
        for table in ttable:
            daytime = getday(self.url)
            print(table[0])
            if table[0] == nad:
                for line in table:
                    if line == nad:
                        continue
                    if len(line) == 7:
                        line.pop(0)
                    nadobjects.append(
                        nadomescanje(
                            ura=line[0],
                            dan=daytime,
                            razred=line[1].replace(". ", ""),
                            ucilnica=line[2],
                            ucitelj=line[3],
                            predmet=line[4],
                        )
                    )
            if table[0] == menur:
                for line in table:
                    if line == menur:
                        continue
                    nadobjects.append(
                        nadomescanje(
                            ura=line[1],
                            dan=daytime,
                            razred=line[0].replace(". ", ""),
                            ucilnica=line[4],
                            ucitelj=line[2].split("→")[1],
                            predmet=line[3].split("→")[1],
                        )
                    )
        return nadobjects


class nadomescanja:
    def __init__(self, fileloc):
        self.config = None
        with open(fileloc) as yamlfile:
            config = yaml.load(yamlfile, Loader=yaml.FullLoader)
        if config == None:
            raise FileNotFoundError("Yaml file does not exist")
        self.token = config["token"]
        self.url = config["url"]
        self.courseid = config["courseid"]
        self.tokenizeurls = config["tokenizeurls"]
        self.minid = config["minid"]
        self.external_cron = config["external_cron"]

    def saveid(self, ids):
        id_ = db.session.query(Lasttime).get(1)
        if id_ == None:
            db.session.add(Lasttime(time=0))
        else:
            id_.time = ids
        db.session.commit()

    def loadid(self):
        id_ = db.session.query(Lasttime).get(1)
        if id_ == None:
            db.session.add(Lasttime(time=0))
            db.session.commit()
            print("0 added", id_)
            return 0
        return id_.time

    def treaturl(self, url):
        for tokenurl in self.tokenizeurls:
            if tokenurl in url:
                if "?" in url:
                    url += "&"
                else:
                    url += "?"
                url += "token=" + self.token
                return url
        return url

    def getupdates(self):
        params = (("moodlewsrestformat", "json"),)
        data = {
            "wstoken": self.token,
            "wsfunction": "core_course_get_contents",
            "courseid": str(self.courseid),
        }
        response = requests.post(self.url, params=params, data=data,)
        resjs = response.json()
        mod = []
        for x in resjs:
            mod += x["modules"]
        lastid = self.loadid()
        maxid = lastid
        fileurls = []
        for modul in mod:
            if "contents" not in modul.keys():
                continue
            fileurl = self.treaturl(modul["contents"][0]["fileurl"])
            fileurls.append(fileurl)
        return fileurls


ndg = nadomescanja(dirname + "/nadconfig.yaml")


# running schedular
def procces_urls():
    debug("Checking data from e-classroom")
    files_updated = False

    nad_url_part = "www.dropbox.com"
    url_list = [
        x.replace("dl=0", "dl=1") for x in ndg.getupdates() if nad_url_part in x
    ]

    if len(url_list) == 0:
        return

    for url in url_list:
        urlfile = PdfFile(url)
        h = urlfile.download()
        u = db.session.query(Urlstring).filter_by(dan=getday(url)).first()
        print(h)
        if u:
            if not u.file_hash.strip() == h:
                debug("Updating existing file")
                files_updated = True

                u.url = url
                u.dan = getday(url)
                u.file_hash = h
                db.session.commit()
            else:
                continue
        else:
            debug("Adding new file")
            files_updated = True

            db.session.add(Urlstring(url=url, dan=getday(url), file_hash=h))

        objcts = urlfile.parsepdf()
        for ele in objcts:
            c = (
                db.session.query(nadomescanje)
                .filter_by(ura=ele.ura, dan=ele.dan, razred=ele.razred)
                .first()
            )

            if c:
                c.predmet = ele.predmet
                c.ucilnica = ele.ucilnica
                c.ucitelj = ele.ucitelj
            else:
                db.session.add(ele)

    if not files_updated:
        debug("No new data found")

    db.session.commit()


if not ndg.external_cron:
    cron = BackgroundScheduler(daemon=True)
    cron.add_job(procces_urls, "cron", minute="10")
    cron.start()

    # Shutdown your cron thread if the web process is stopped
    atexit.register(lambda: cron.shutdown(wait=False))


@app.route("/nad")
def getnad():
    if request.args.get("razred"):
        qry = (
            db.session.query(nadomescanje)
            .filter_by(dan=request.args.get("dan"), razred=request.args.get("razred"))
            .all()
        )
    else:
        qry = (
            db.session.query(nadomescanje).filter_by(dan=request.args.get("dan")).all()
        )
    json_lst = [x.get_dict() for x in qry]
    return jsonify(json_lst)


@app.route("/url")
def geturl():
    qry = db.session.query(Urlstring).filter_by(dan=request.args.get("dan")).all()
    json_lst = [x.get_dict() for x in qry]
    return jsonify(json_lst)


@app.after_request
def apply_caching(response):
    response.headers["Access-Control-Allow-Origin"] = "https://gimvicurnik.filips.si"
    return response


if __name__ == "__main__":
    if not ndg.external_cron:
        procces_urls()
    app.run()
